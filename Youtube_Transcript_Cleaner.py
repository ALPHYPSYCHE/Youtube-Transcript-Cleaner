import re
import os
import tkinter as tk
from tkinter import messagebox
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx import Document
import win32com.client as win32

# --- Logic Functions ---

def convert_doc_to_docx(doc_path):
    """Converts legacy .doc to .docx on Windows"""
    try:
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        docx_path = os.path.splitext(doc_path)[0] + "_temp.docx"
        doc.SaveAs2(docx_path, FileFormat=16)
        doc.Close()
        return docx_path
    except Exception as e:
        print(f"Error converting .doc: {e}")
        return None

def is_garbage_line(text):
    """Regex pattern to identify timestamps and durations"""
    if re.match(r'^\d+:\d+(?:\:\d+)?$', text):
        return True
    if re.match(r'^(\d+\s+(second|minute|hour)s?)(?:\s*,\s*\d+\s+(second|minute|hour)s?)*$', text, re.IGNORECASE):
        return True
    return False

def process_file(input_path):
    """Main processing core that detects extension and cleans the file"""
    input_path = input_path.strip('{}""\'\'')
    
    if not os.path.exists(input_path):
        return False, f"File not found:\n{input_path}"

    input_dir = os.path.dirname(input_path)
    file_name = os.path.basename(input_path)
    base_name, ext = os.path.splitext(file_name)
    ext = ext.lower()
    
    temp_docx_created = False
    actual_input = input_path

    # 1. Handle .doc format
    if ext == '.doc':
        actual_input = convert_doc_to_docx(input_path)
        if not actual_input:
            return False, "Error converting legacy .doc file."
        temp_docx_created = True
        ext = '.docx'

    # 2. Process Word (.docx) format
    if ext == '.docx':
        final_output = os.path.join(input_dir, f"CLEANED_{base_name}.docx")
        try:
            doc = Document(actual_input)
            new_doc = Document()
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and not is_garbage_line(text):
                    new_doc.add_paragraph(text)
            new_doc.save(final_output)
        except Exception as e:
            return False, f"Error processing DOCX: {e}"
        finally:
            if temp_docx_created and os.path.exists(actual_input):
                os.remove(actual_input)

    # 3. Process Text (.txt) format
    elif ext == '.txt':
        final_output = os.path.join(input_dir, f"CLEANED_{base_name}.txt")
        try:
            with open(input_path, 'r', encoding='utf-8') as infile, open(final_output, 'w', encoding='utf-8') as outfile:
                for line in infile:
                    text = line.strip()
                    if text and not is_garbage_line(text):
                        outfile.write(text + "\n")
        except Exception as e:
            return False, f"Error processing TXT: {e}"
    else:
        return False, f"Unsupported file format: {ext}"

    return True, final_output

# --- GUI Actions ---

def drop_file(event):
    """Triggers instantly when a file is dropped"""
    file_path = event.data.strip('{}""\'\'')
    
    # Update interface to STARTED
    status_label.config(text="STATUS: STARTED", fg="#e65100")
    root.update()
    
    # Run the core cleaner function
    success, message = process_file(file_path)
    
    if success:
        status_label.config(text="STATUS: FINISHED", fg="#2e7d32")
        messagebox.showinfo("Success", f"File processed successfully!\n\nSaved as:\n{message}")
    else:
        status_label.config(text="STATUS: FAILED", fg="#d32f2f")
        messagebox.showerror("Error", message)

# --- GUI Layout Building ---

root = TkinterDnD.Tk()
root.title("Instant Transcript Cleaner [By Alphypsyche]")
root.geometry("500x260")
root.resizable(False, False)

# Description Title
lbl_title = tk.Label(root, text="Instant Subtitle & Transcript Cleaner", font=("Arial", 12, "bold"))
lbl_title.pack(pady=(20, 5))

lbl_desc = tk.Label(root, text="Drop your file below. It will clean and save to the same folder automatically.", font=("Arial", 9), fg="#555555")
lbl_desc.pack(pady=(0, 15))

# Drag & Drop Box Area
drop_zone = tk.Label(
    root, 
    text="\n【 Drop File Here to Start 】\n\nSupports: .txt, .docx, .doc", 
    bg="#f0f4c3", 
    fg="#33691e", 
    bd=2, 
    relief="groove", 
    font=("Arial", 11, "bold")
)
drop_zone.pack(fill=tk.X, padx=40, pady=5)

# Register Drop Trigger
drop_zone.drop_target_register(DND_FILES)
drop_zone.dnd_bind('<<Drop>>', drop_file)

# Real-time Status Tracker Bar
status_label = tk.Label(root, text="STATUS: IDLE", font=("Arial", 10, "bold"), fg="#757575")
status_label.pack(pady=(20, 0))

root.mainloop()
